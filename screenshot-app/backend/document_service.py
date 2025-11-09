"""
Document Service
Generates Word documents with screenshots
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
from typing import List
from pathlib import Path
from datetime import datetime

class DocumentService:
    def __init__(self):
        self.max_image_width = 6.0  # inches
    
    async def generate(
        self,
        screenshot_paths: List[str],
        output_path: str,
        title: str = "Screenshot Report"
    ) -> str:
        """
        Generate Word document with screenshots
        
        Args:
            screenshot_paths: List of screenshot file paths
            output_path: Output .docx file path
            title: Document title
        
        Returns:
            Path to generated document
        """
        # Create document
        doc = Document()
        
        # Add title
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(title)
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add generation date
        date_paragraph = doc.add_paragraph()
        date_run = date_paragraph.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        date_run.font.size = Pt(10)
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacing
        
        # Add screenshots
        for i, screenshot_path in enumerate(screenshot_paths, 1):
            if not Path(screenshot_path).exists():
                continue

            # Get filename and extract clean name
            filename = Path(screenshot_path).name

            # Extract clean name from filename (remove segment numbers and extension)
            # Example: Accounting_AutoPostingSettings_001.png -> Accounting_AutoPostingSettings
            clean_name = filename.replace('.png', '').replace('.jpg', '').replace('.jpeg', '')

            import re

            # Check if this is a timestamp-based filename (e.g., screenshot_1_1762613242452)
            # Pattern: screenshot_N_TIMESTAMP or domain_TIMESTAMP
            if re.match(r'^screenshot_\d+_\d+$', clean_name):
                # Timestamp-based: Use "Screenshot N" format
                match = re.match(r'^screenshot_(\d+)_\d+$', clean_name)
                if match:
                    display_name = f"Screenshot {match.group(1)}"
                else:
                    display_name = f"Screenshot {i}"
            elif re.match(r'^.+_\d{8}_\d{6}(_\d{3})?$', clean_name):
                # Domain + timestamp format (e.g., example.com_20251108_201937_001)
                # Extract domain part
                domain_part = re.sub(r'_\d{8}_\d{6}(_\d{3})?$', '', clean_name)
                display_name = domain_part.replace('_', ' ').title()
            else:
                # PascalCase format (e.g., Accounting_AutoPostingSettings_001)
                # Remove segment numbers (_001, _002, etc.)
                clean_name = re.sub(r'_\d{3}$', '', clean_name)
                # Convert underscores to spaces for better readability
                display_name = clean_name.replace('_', ' ')

            # Add heading with clean name
            heading = doc.add_heading(display_name, level=2)

            # Add filename as caption
            doc.add_paragraph(f"File: {filename}", style='Caption')
            
            # Process and add image
            try:
                # Get image dimensions
                img = Image.open(screenshot_path)
                img_width, img_height = img.size
                
                # Calculate display width (maintain aspect ratio)
                aspect_ratio = img_height / img_width
                display_width = min(self.max_image_width, img_width / 96)  # 96 DPI
                display_height = display_width * aspect_ratio
                
                # Add image to document
                doc.add_picture(screenshot_path, width=Inches(display_width))
                
                # Add spacing
                doc.add_paragraph()
                
            except Exception as e:
                doc.add_paragraph(f"Error adding image: {str(e)}", style='Caption')
        
        # Save document
        # Expand ~ to home directory
        output_path = Path(output_path).expanduser()
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))

        return str(output_path)

